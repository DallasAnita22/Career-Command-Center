import io
import re
from docx import Document
from fpdf import FPDF
import fitz  # PyMuPDF

def extract_text_from_file(uploaded_file):
    text = ""
    try:
        if uploaded_file.name.endswith(".pdf"):
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                for page in doc: text += page.get_text()
        elif uploaded_file.name.endswith(".docx"):
            doc = Document(uploaded_file)
            for para in doc.paragraphs: text += para.text + "\n"
        else:
            text = uploaded_file.read().decode("utf-8")
    except Exception as e: return f"Error reading file: {str(e)}"
    return text

def sanitize_text(text):
    if not text: return ""
    replacements = {"\u2013": "-", "\u2014": "--", "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"', "\u2022": "-", "…": "..."}
    for char, r in replacements.items(): text = text.replace(char, r)
    return text.encode('latin-1', 'replace').decode('latin-1')

def generate_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    name = sanitize_text(data.get("name", ""))
    contact = sanitize_text(f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('linkedin', '')}")
    
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, name, ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, contact, ln=True, align="C")
    pdf.ln(10)
    
    sections = [("PROFESSIONAL SUMMARY", "summary"), ("EXPERIENCE", "experience"), ("SKILLS & EDUCATION", "skills"), ("REFERENCES", "references")]
    for title, key in sections:
        if data.get(key):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, title, ln=True, border="B")
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, sanitize_text(data[key]))
            pdf.ln(5)
    return bytes(pdf.output())

def generate_docx(data):
    doc = Document()
    doc.add_heading(data.get("name", ""), 0)
    doc.add_paragraph(f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('linkedin', '')}")
    
    sections = [('Professional Summary', 'summary'), ('Experience', 'experience'), ('Skills & Education', 'skills'), ('References', 'references')]
    for title, key in sections:
        if data.get(key):
            doc.add_heading(title, level=1)
            doc.add_paragraph(data[key])
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def parse_resume_to_dict(text):
    data = {"name": "", "email": "", "phone": "", "linkedin": "", "summary": "", "experience": "", "skills": "", "references": ""}
    if not text: return data
    
    email = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if email: data["email"] = email.group(0)
    
    phone = re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    if phone: data["phone"] = phone.group(0)
    
    link = re.search(r'(https?://[^\s]+)', text)
    if link: data["linkedin"] = link.group(0)
    
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if lines and "resume" not in lines[0].lower(): data["name"] = lines[0]
    
    lower = text.lower()
    idxs = {"exp": lower.find("experience"), "skl": lower.find("skills"), "edu": lower.find("education"), "ref": lower.find("references")}
    
    sorted_idxs = sorted([i for i in idxs.values() if i != -1])
    if sorted_idxs:
        data["summary"] = text[:sorted_idxs[0]].replace(data["name"], "").replace(data["email"], "").replace(data["phone"], "").strip()
    
    if idxs["exp"] != -1:
        end = min([i for i in sorted_idxs if i > idxs["exp"]], default=len(text))
        data["experience"] = text[idxs["exp"]:end].replace("Experience", "", 1).strip()
        
    if idxs["skl"] != -1:
        end = min([i for i in sorted_idxs if i > idxs["skl"]], default=len(text))
        data["skills"] = text[idxs["skl"]:end].replace("Skills", "", 1).strip()
        
    if idxs["ref"] != -1:
        data["references"] = text[idxs["ref"]:].replace("References", "", 1).strip()
        
    return data